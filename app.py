import os
import tempfile
import json
import spacy
from flask import Flask, render_template, request, send_file
from docx import Document
from collections import Counter
import google.generativeai as genai

# Use Gemini API key from environment variable
GEMINI_API_KEY = os.getenv("m")
genai.configure(api_key=GEMINI_API_KEY)

app = Flask(__name__)
nlp = spacy.load("en_core_web_sm")

def extract_keywords(job_text, top_n=15):
    doc = nlp(job_text.lower())
    keywords = [token.lemma_ for token in doc if token.pos_ in ['NOUN', 'PROPN', 'VERB', 'ADJ'] and not token.is_stop]
    most_common = [word for word, count in Counter(keywords).most_common(top_n)]
    return set(most_common)

def match_and_score_section(text, keywords):
    text = text.lower()
    return sum(1 for kw in keywords if kw in text)

def prioritize_experience(cv, keywords):
    scored = []
    for exp in cv.get("experience", []):
        section_text = " ".join(exp.get("description", [])) if isinstance(exp.get("description"), list) else exp.get("description", "")
        score = match_and_score_section(section_text, keywords)
        scored.append((score, exp))
    scored.sort(reverse=True, key=lambda x: x[0])
    return [exp for score, exp in scored]

def prioritize_skills(skills, keywords):
    all_skills = []
    if isinstance(skills, dict):
        for v in skills.values():
            all_skills.extend(v)
    elif isinstance(skills, list):
        all_skills = skills
    return sorted(set(all_skills), key=lambda s: any(k in s.lower() for k in keywords), reverse=True)

def gemini_rewrite_bullet(bullet, job_desc):
    prompt = f"""
    You are an expert ATS resume writer. Here is an original resume bullet: "{bullet}"
    Here is the target job description: "{job_desc}"
    Thoroughly rewrite the bullet so it uses exact language and keywords from the job description. Make it concise, results-oriented, and perfectly tailored for ATS matching.
    """
    try:
        model = genai.GenerativeModel('gemini-1.5-pro')
        response = model.generate_content(prompt)
        text = response.text.strip()
        return text if text else bullet
    except Exception as e:
        print(f"Gemini error: {e}")
        return bullet

def generate_docx(cv, job_keywords, job_desc):
    doc = Document()
    doc.add_heading(cv.get("name", ""), 0)
    contact = cv.get("contact", {})
    doc.add_paragraph(
        f"{contact.get('location','')} | {contact.get('phone','')} | {contact.get('email','')}\n"
        f"{contact.get('linkedin','')}"
    )
    if "summary" in cv:
        doc.add_heading("Professional Summary", level=1)
        doc.add_paragraph(cv["summary"])

    doc.add_heading("Skills", level=1)
    skills = prioritize_skills(cv.get("skills", []), job_keywords)
    if skills:
        doc.add_paragraph(", ".join(skills))

    doc.add_heading("Experience", level=1)
    experiences = prioritize_experience(cv, job_keywords)
    for exp in experiences:
        title = exp.get("title", "")
        company = exp.get("company", "")
        dates = exp.get("dates", "")
        location = exp.get("location", "")
        doc.add_heading(f"{title} – {company}", level=2)
        doc.add_paragraph(f"{dates} | {location}")
        desc = exp.get("description", [])
        if isinstance(desc, list):
            for bullet in desc:
                tailored_bullet = gemini_rewrite_bullet(bullet, job_desc)
                doc.add_paragraph(tailored_bullet, style="List Bullet")
        else:
            doc.add_paragraph(desc)

    if cv.get("projects"):
        doc.add_heading("Projects", level=1)
        for proj in cv["projects"]:
            doc.add_heading(proj.get("name", ""), level=2)
            doc.add_paragraph(proj.get("description", ""))

    if cv.get("certifications"):
        doc.add_heading("Certifications", level=1)
        for cert in cv["certifications"]:
            doc.add_paragraph(cert)

    if cv.get("education"):
        doc.add_heading("Education", level=1)
        for edu in cv["education"]:
            deg = edu.get("degree", "")
            inst = edu.get("institution", "")
            dates = edu.get("dates", "")
            cgpa = edu.get("cgpa", "")
            edu_line = f"{deg}, {inst} ({dates})"
            if cgpa:
                edu_line += f", CGPA: {cgpa}"
            doc.add_paragraph(edu_line)

    temp = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp.name)
    return temp.name

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        # Validation
        cv_file = request.files.get('cv_file')
        if not cv_file or cv_file.filename == '':
            return render_template('index.html', error="Please upload your master CV in JSON format.")
        if not cv_file.filename.lower().endswith('.json'):
            return render_template('index.html', error="Only .json files are accepted for your resume.")
        try:
            cv_data = json.load(cv_file)
        except Exception as e:
            return render_template('index.html', error="Invalid JSON file. Please upload a valid resume file.")

        job_desc = request.form.get('job_desc', '')
        if not job_desc or not job_desc.strip():
            return render_template('index.html', error="Please paste the job description.")

        keywords = extract_keywords(job_desc)
        docx_path = generate_docx(cv_data, keywords, job_desc)
        return send_file(
            docx_path,
            as_attachment=True,
            download_name='ATS_CV.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    return render_template('index.html')

if __name__ == '__main__':
    app.run(host="0.0.0.0", port=5000, debug=True)
